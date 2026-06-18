"""
VerdaSense FYP2 — Supervisor Meeting Proposal Document Generator
Output: docs/VerdaSense_FYP2_Supervisor_Proposal.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\GIGA\OneDrive - Universiti Malaya\Documents\rag-for-beginners"
OUT  = os.path.join(BASE, "docs", "VerdaSense_FYP2_Supervisor_Proposal.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

# ── Colour palette ─────────────────────────────────────────────────────────────
C_DARK_BLUE  = "1F3864"
C_MED_BLUE   = "2E74B5"
C_LIGHT_BLUE = "EAF0FB"
C_TEAL       = "1B6B6B"
C_GREEN      = "375623"
C_GREEN_BG   = "E2EFDA"
C_ORANGE     = "C55A11"
C_ORANGE_BG  = "FCE4D6"
C_RED        = "C00000"
C_RED_BG     = "FFDEDE"
C_GREY_BG   = "F2F2F2"
C_YELLOW_BG  = "FFFEF0"
C_WHITE      = "FFFFFF"

# ── XML helpers ────────────────────────────────────────────────────────────────
def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def cell_border(cell, color="AAAAAA", sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcB.append(b)
    tcPr.append(tcB)

def para_run(cell_or_doc, text, bold=False, italic=False,
             size=10, color="000000", align=None, space_before=0, space_after=2):
    p = cell_or_doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p

def shade_para(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def section_heading(doc, text, fill=C_DARK_BLUE, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f"  {text}  ")
    r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_para(p, fill)
    return p

def sub_heading(doc, text, color=C_MED_BLUE, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p

def bullet(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

def info_box(doc, text, fill=C_LIGHT_BLUE, border_color="2E74B5", size=9):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    c = tbl.rows[0].cells[0]
    shade_cell(c, fill)
    cell_border(c, border_color, "8")
    c.paragraphs[0].clear()
    p = c.add_paragraph(text)
    p.runs[0].font.size = Pt(size)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def two_col_table(doc, rows_data, col_widths=(5, 12),
                  hdr_fill=C_MED_BLUE, row_fill=C_LIGHT_BLUE, alt_fill=C_WHITE):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(col_widths[0])
    tbl.columns[1].width = Cm(col_widths[1])
    for i, (label, value) in enumerate(rows_data):
        row = tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        fill = row_fill if i % 2 == 0 else alt_fill
        shade_cell(lc, hdr_fill if label.startswith("__HDR__") else fill)
        shade_cell(vc, hdr_fill if label.startswith("__HDR__") else fill)
        cell_border(lc); cell_border(vc)
        lc.paragraphs[0].clear(); vc.paragraphs[0].clear()
        lbl = label.replace("__HDR__", "")
        lr = lc.add_paragraph(lbl).runs
        if lr: lr[0].bold = True; lr[0].font.size = Pt(9)
        if label.startswith("__HDR__"):
            lr[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        else:
            lr[0].font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
        vp = vc.add_paragraph(value)
        vp.runs[0].font.size = Pt(9)
        if label.startswith("__HDR__"):
            vp.runs[0].bold = True
            vp.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def status_badge(doc, items):
    """items = list of (label, fill, text_color)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    for label, _, _ in items:
        r = p.add_run(f"  {label}  ")
        r.bold = True; r.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.2)

    # ── COVER ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VerdaSense RAG System")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(C_DARK_BLUE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("FYP2 Proposal — Supervisor Meeting")
    r2.bold = True; r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor.from_string(C_MED_BLUE)

    doc.add_paragraph()
    meta = [
        ("Student",  "Tee Qi Jing (23004894)"),
        ("Project",  "Evidence-Grounded Clinical Decision Support for Wound Dressing Recommendation"),
        ("Meeting",  "FYP2 Supervisor Meeting — June 2026"),
        ("Prepared", "Tee Qi Jing · Universiti Malaya"),
    ]
    for label, val in meta:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp.add_run(f"{label}:  "); mr.bold = True; mr.font.size = Pt(11)
        mr2 = mp.add_run(val); mr2.font.size = Pt(11)

    doc.add_paragraph()
    agenda_items = [
        ("1", "FYP1 Recap — What Was Built & Best Results"),
        ("2", "FYP1 Viva Feedback & Clinical Collaborator Feedback"),
        ("3", "FYP2 Core Contribution: Multimodal RAG (G4)"),
        ("4", "FYP2 Supporting Changes: UI/UX + Product Gallery"),
        ("5", "FYP2 Secondary: Wound Category Classification"),
        ("6", "Dropped: Conversational RAG (rationale)"),
        ("7", "FYP2 Timeline (4–5 months)"),
        ("8", "Risks & Mitigations"),
        ("9", "Questions for Supervisor"),
    ]
    at = doc.add_table(rows=0, cols=2)
    at.style = "Table Grid"
    at.columns[0].width = Cm(1.5)
    at.columns[1].width = Cm(14)
    hr = at.add_row()
    shade_cell(hr.cells[0], C_DARK_BLUE); shade_cell(hr.cells[1], C_DARK_BLUE)
    cell_border(hr.cells[0]); cell_border(hr.cells[1])
    hr.cells[0].paragraphs[0].clear(); hr.cells[1].paragraphs[0].clear()
    r = hr.cells[0].add_paragraph("#").runs[0]
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(255,255,255)
    r = hr.cells[1].add_paragraph("Agenda").runs[0]
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(255,255,255)
    for num, item in agenda_items:
        row = at.add_row()
        shade_cell(row.cells[0], C_LIGHT_BLUE)
        shade_cell(row.cells[1], C_WHITE)
        cell_border(row.cells[0]); cell_border(row.cells[1])
        row.cells[0].paragraphs[0].clear(); row.cells[1].paragraphs[0].clear()
        r = row.cells[0].add_paragraph(num).runs[0]
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(C_MED_BLUE)
        r = row.cells[1].add_paragraph(item).runs[0]; r.font.size = Pt(10)

    doc.add_page_break()

    # ── SECTION 1: FYP1 RECAP ─────────────────────────────────────────────────
    section_heading(doc, "1.  FYP1 Recap — What Was Built")

    sub_heading(doc, "System: VerdaSense Unimodal RAG")
    info_box(doc,
        "VerdaSense is a single-turn, evidence-grounded wound dressing recommendation system. "
        "It receives T.I.M.E. wound assessment inputs (Tissue, Infection, Moisture, Edge) from a "
        "companion CV pipeline, retrieves relevant chunks from an 8-source clinical knowledge base "
        "via dense semantic search, and generates a structured 9-section recommendation using a "
        "grounded system prompt. Every clinical claim is cited back to a specific guideline source.",
        fill=C_LIGHT_BLUE)

    sub_heading(doc, "Knowledge Base & Evaluation Infrastructure")
    two_col_table(doc, [
        ("__HDR__Component", "__HDR__Detail"),
        ("KB Sources",       "8 clinical guidelines (WCM, GP-MOH, AJGP, SFP, EWMA, ISTAP, ANZBA, RCH)"),
        ("KB Chunks",        "138 manually curated chunks with metadata (wound_type, wound_category, authority, year)"),
        ("Embedding DB",     "ChromaDB — db_wound_care_v4_bge/ — BAAI/bge-large-en-v1.5 (BGE Large)"),
        ("Testset",          "32 manually curated cases — 5 categories (Cat A–E) — wound_testset_v3.json"),
        ("RAGAS Judge",      "gpt-4o-mini + text-embedding-3-small (fixed across all experiments)"),
        ("Ablation runs",    "R1–R5 (retrieval) + G1–G3 (generation) — 8 experiments, 3 runs each"),
    ])

    sub_heading(doc, "Ablation-Best Configuration (Current System)")
    two_col_table(doc, [
        ("__HDR__Experiment", "__HDR__Winner & Setting"),
        ("R1 — Query Strategy",    "R1-C: Multi-axis sub-queries (A: wound type pin, B: mechanism, C: notes)"),
        ("R2 — Retrieval Method",  "R2-A: Dense-only (no BM25, no reranker)"),
        ("R3 — Top-K Depth",       "R3-C: k = 6"),
        ("R4 — Embedding Model",   "R4-B: BAAI/bge-large-en-v1.5 (BGE Large)"),
        ("R5 — Multimodal (new)",  "R5-A baseline reproduced; caption injection mixed results — extends to G4"),
        ("G1 — Prompt Strategy",   "G1-C: Grounded system prompt with strict citation rules"),
        ("G2 — Closed-Source LLM", "G2-D: Gemini 2.5 Flash — FA=0.8147, Safety=90.6%"),
        ("G3 — Open-Source LLM",   "G3-A: Qwen3:14b (HPC A100) — FA=0.73, Safety=80.2% (no open model passes gate)"),
    ])

    sub_heading(doc, "Best Evaluation Results (G2-D: Gemini 2.5 Flash)")
    two_col_table(doc, [
        ("__HDR__Metric",           "__HDR__Score"),
        ("Faithfulness (FA)",        "0.8147 ± 0.0100  (81.5% of claims grounded in retrieved evidence)"),
        ("Answer Relevancy (AR)",    "0.6770 ± 0.0210"),
        ("Context Recall (CR)",      "0.8945  (89.5% of needed chunks retrieved)"),
        ("Safety Pass Rate",         "90.6%  (29/32 cases pass all clinical safety rules)"),
        ("Classifier Accuracy",      "87.5%  (28/32 correct wound type classification)"),
        ("Generation Latency",       "~18 s avg (Gemini 2.5 Flash API)"),
    ], col_widths=(5.5, 11.5))

    doc.add_page_break()

    # ── SECTION 2: FEEDBACK ───────────────────────────────────────────────────
    section_heading(doc, "2.  Feedback — FYP1 Viva Panel + Clinical Collaborator (Ms Saw)")

    sub_heading(doc, "FYP1 Viva Panel Feedback")
    two_col_table(doc, [
        ("__HDR__Panel Question",    "__HDR__FYP2 Response"),
        ("Why not rule-based only?",
         "VerdaSense is already hybrid: classify_wound() is pure rule-based (WT1-8 + referral/ABx logic). "
         "RAG handles evidence grounding, dressing rationale, citation, and etiology nuance that rules cannot. "
         "FYP2 strengthens this narrative — no new experiment needed."),
        ("How do you know output is clinically accurate?",
         "FYP2 adds clinician human evaluation: Ms Saw reviews 8 cases, rates clinical concordance (Likert 1-5). "
         "Closes the 'automated metrics only' gap."),
        ("Can it handle images?",
         "FYP2 primary contribution: G4 Multimodal RAG — wound images processed by VLLM, "
         "clinical captions injected into retrieval and generation. Directly addresses this gap."),
    ], col_widths=(5, 12))

    doc.add_paragraph()
    sub_heading(doc, "Clinical Collaborator Feedback (Ms Saw Shier Khee, Surgeon)")
    two_col_table(doc, [
        ("__HDR__Ms Saw's Feedback",  "__HDR__FYP2 Action"),
        ("'That's why need photo' — T.I.M.E. labels alone can't determine referral for infected wounds",
         "G4 Multimodal RAG: add wound image as input, VLLM generates clinical caption to enrich recommendation"),
        ("Remove 'Rationale by T.I.M.E. Factor' section — not useful for patients",
         "UI redesign: remove this section from patient-facing output"),
        ("Add description explaining primary & secondary dressing + show dressing images",
         "Product Gallery + dressing type descriptions with images in Tab 1 result"),
        ("Output too texty — needs to be short, concise, least clinical terminology for elderly patients",
         "Prompt redesign for FYP2: patient-friendly language, shorter structured output"),
        ("Show dressing product images (hydrocolloid, silver) in primary/secondary dressing",
         "Product Gallery: map recommendation → Big Pharmacy product images (157 products scraped)"),
        ("KB sources: RCH is paediatric-only — not appropriate for adult patients",
         "RCH metadata fix: add population=paediatric tag, exclude from adult patient retrieval"),
    ], col_widths=(7, 10))

    doc.add_page_break()

    # ── SECTION 3: MULTIMODAL RAG ─────────────────────────────────────────────
    section_heading(doc, "3.  FYP2 Core Contribution: Multimodal RAG (G4 Ablation)")

    info_box(doc,
        "PRIMARY FYP2 RESEARCH CONTRIBUTION\n"
        "Research question: Does injecting VLLM-generated visual wound descriptions into a text-only RAG "
        "pipeline improve evidence-grounded dressing recommendation quality, and where is the best injection point?\n"
        "This directly responds to Ms Saw's clinical feedback that wound images contain information "
        "that structured T.I.M.E. labels alone cannot capture.",
        fill=C_GREEN_BG, border_color=C_GREEN)

    sub_heading(doc, "FYP1 vs FYP2 Architecture")

    arc_tbl = doc.add_table(rows=0, cols=2)
    arc_tbl.style = "Table Grid"
    arc_tbl.columns[0].width = Cm(8.5)
    arc_tbl.columns[1].width = Cm(8.5)

    def arc_hdr(cells, t1, t2, fill):
        for c, t in zip(cells, [t1, t2]):
            shade_cell(c, fill); cell_border(c)
            c.paragraphs[0].clear()
            r = c.add_paragraph(t).runs[0]
            r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255,255,255)

    def arc_row(cells, t1, t2, f1=C_GREY_BG, f2=C_YELLOW_BG):
        for c, t, f in zip(cells, [t1, t2], [f1, f2]):
            shade_cell(c, f); cell_border(c)
            c.paragraphs[0].clear()
            r = c.add_paragraph(t).runs[0]; r.font.size = Pt(9)

    hr = arc_tbl.add_row()
    arc_hdr(hr.cells, "FYP1 — Unimodal (Current System)", "FYP2 — Multimodal (Proposed)", C_MED_BLUE)

    rows_arc = [
        ("INPUT\nStructured T.I.M.E. payload only\n(Tissue %, Infection, Moisture, Edge, Notes)",
         "INPUT\nStructured T.I.M.E. payload + Wound Image\n(same T.I.M.E. labels + raw photograph)"),
        ("STEP 1: classify_wound() [Rule-based]\n"
         "Assigns WT1-8, referral, ABx, etiology",
         "STEP 1: classify_wound() [Rule-based] — UNCHANGED\n"
         "Assigns WT1-8, referral, ABx, etiology"),
        ("STEP 2: Multi-axis Retrieval (R1-C)\n"
         "Sub-A: Wound type algorithm chunk\n"
         "Sub-B: Dressing mechanism query\n"
         "Sub-C: Patient notes query\n"
         "Top-K=6 dense chunks from ChromaDB",
         "STEP 2: Multi-axis Retrieval (R1-C ENHANCED)\n"
         "Sub-A: Wound type algorithm chunk — UNCHANGED\n"
         "Sub-B: Dressing mechanism query — UNCHANGED\n"
         "Sub-C: Patient notes + VLLM caption (ENRICHED)\n"
         "[Optional Sub-D: Caption-only axis]\n"
         "Top-K=6 dense chunks from ChromaDB"),
        ("STEP 3: Generation (G1-C Grounded Prompt)\n"
         "assessment_text = T.I.M.E. labels only\n"
         "9-section structured output\n"
         "Cites sources 1-6",
         "STEP 3: Generation (G4 Grounded Prompt)\n"
         "assessment_text = T.I.M.E. + VLLM visual description (NEW)\n"
         "9-section output (patient-friendly language)\n"
         "Cites sources 1-6"),
        ("OUTPUT\n9 clinical sections (GP-facing language)\n"
         "No dressing images\nNo product gallery",
         "OUTPUT (UPDATED)\n"
         "Simplified sections (patient-facing language)\n"
         "Dressing images + product gallery\n"
         "Evidence sources retained"),
    ]
    for r1, r2 in rows_arc:
        row = arc_tbl.add_row()
        arc_row(row.cells, r1, r2)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    sub_heading(doc, "VLLM Caption Generation")
    two_col_table(doc, [
        ("__HDR__Component", "__HDR__Detail"),
        ("VLLM Model",          "Gemini 2.5 Flash (Vision) — same model as G2-D generation winner"),
        ("Caption format",      "Structured 4-axis: TISSUE / INFECTION / MOISTURE / EDGE + ADDITIONAL"),
        ("Caption latency",     "~8–13 s per image (OpenRouter API, R5 benchmarks)"),
        ("Injection point 1",   "Sub-query C: patient notes + VLLM caption combined as retrieval query"),
        ("Injection point 2",   "assessment_text in generation: T.I.M.E. labels + 'VISUAL WOUND ANALYSIS' section"),
        ("Known limitation",    "VLLM may miss early infection signs not visually obvious (WT03/04 in R5). "
                                "T.I.M.E. infection label retains priority for ABx/referral decisions."),
    ])

    sub_heading(doc, "Caption-T.I.M.E. Conflict Handling")
    info_box(doc,
        "When VLLM caption contradicts T.I.M.E. labels (e.g., caption says 'no infection' but T.I.M.E. "
        "says 'locally infected'):\n"
        "  Rule 1 — T.I.M.E. Infection axis takes priority: ABx and referral decisions follow T.I.M.E. labels, "
        "not caption.\n"
        "  Rule 2 — VLLM caption enriches tissue description and moisture detail where T.I.M.E. labels are coarse.\n"
        "  Rule 3 — Discrepancy is logged as a retrieval note for transparency.\n"
        "  Future work: learned fusion layer to weight T.I.M.E. vs caption per axis.",
        fill=C_ORANGE_BG, border_color=C_ORANGE)

    sub_heading(doc, "G4 Ablation Design")
    two_col_table(doc, [
        ("__HDR__Config",    "__HDR__What it tests"),
        ("G4-A (Baseline)",  "Text-only T.I.M.E. — same as G2-D winner. Cross-check against FYP1 results."),
        ("G4-B",             "VLLM caption injected into Sub-query C only (retrieval enrichment)"),
        ("G4-C",             "VLLM caption injected into assessment_text only (generation enrichment)"),
        ("G4-D (Main result)","VLLM caption in BOTH retrieval Sub-C AND assessment_text — combined effect"),
        ("G4-E (Optional)",  "G4-D config but with GPT-4o as VLLM instead of Gemini — compares VLLM model quality"),
    ])

    sub_heading(doc, "G4 Evaluation Metrics")
    two_col_table(doc, [
        ("__HDR__Metric",               "__HDR__What it measures"),
        ("Faithfulness (FA)",            "Are generated claims supported by retrieved chunks? (RAGAS, same judge as FYP1)"),
        ("Answer Relevancy (AR)",        "Does the answer address the wound case? (RAGAS)"),
        ("Safety Pass Rate",             "Does output pass all clinical safety rules? (deterministic checker)"),
        ("Caption Accuracy Rate",        "% captions Ms Saw rates as clinically accurate (from G4 review form)"),
        ("Infection Detection Rate",     "% infected wound types (WT03/04/07/08) where caption correctly flags infection"),
        ("Clinical Concordance Rate",    "% of G4-D recommendations that agree with Ms Saw's recommendations (8 WT cases)"),
        ("Context Recall delta",         "G4-D CR vs G4-A CR — does caption improve chunk retrieval?"),
    ], col_widths=(5.5, 11.5))

    sub_heading(doc, "G4 Experiment Status & Data")
    two_col_table(doc, [
        ("__HDR__Item", "__HDR__Status"),
        ("R5 (retrieval half)",         "COMPLETE — caption injection in retrieval tested, mixed results documented"),
        ("8 WT wound images",           "SELECTED — WT01–WT08, public datasets (Medetec, WSNET)"),
        ("VLLM captions (R5 baseline)", "GENERATED — Qwen-2.5-VL-72B via OpenRouter for all 8 images"),
        ("G4 review form",              "CREATED — VerdaSense_G4_Clinical_Review_Form.docx sent to Ms Saw"),
        ("Ms Saw image review",         "PENDING — awaiting response on image suitability + clinical recommendations"),
        ("G4-A through G4-D notebooks", "TO BUILD in FYP2"),
        ("Conversational testset",      "DROPPED — moved to future work"),
    ])

    doc.add_page_break()

    # ── SECTION 4: UI/UX + PRODUCT GALLERY ───────────────────────────────────
    section_heading(doc, "4.  FYP2 Supporting Changes: UI/UX Redesign + Product Gallery")

    sub_heading(doc, "4.1  UI/UX Redesign (Based on Ms Saw's Feedback)")
    two_col_table(doc, [
        ("__HDR__Change",           "__HDR__Detail"),
        ("Remove section",          "Remove 'Rationale by T.I.M.E. Factor' — Ms Saw: not useful for patients"),
        ("Output language",         "Rewrite generation prompt for patient-facing language: "
                                    "short, concise, no clinical jargon, accessible to elderly users"),
        ("Output structure",        "Reduce from 9 sections to ~5–6 key sections: "
                                    "What to use / What not to use / How to apply / When to change / "
                                    "See a doctor if... / Where to buy"),
        ("Dressing explanation",    "Add brief plain-English description of primary & secondary dressing type"),
        ("Dressing images",         "Show product sample images alongside dressing type name in result"),
    ])

    sub_heading(doc, "4.2  Product Dressing Gallery")
    info_box(doc,
        "Data already available: 157 wound dressing products scraped from Big Pharmacy Malaysia "
        "(bigpharmacy.com.my) using Shopify public JSON API. Each product has: name, vendor, price (MYR), "
        "dressing type classification, product image URL, product page URL, availability status.\n\n"
        "FYP purpose: Show evidence sources and dressing type options in the web interface.\n"
        "Production purpose (Ms Saw): Show actual branded product images with Malaysian prices, "
        "so patients know exactly what to buy at the pharmacy.",
        fill=C_GREEN_BG, border_color=C_GREEN)

    two_col_table(doc, [
        ("__HDR__Dressing Type",     "__HDR__Products available (Big Pharmacy)"),
        ("Silver / Antimicrobial",   "62 products — largest category"),
        ("Gauze",                    "36 products"),
        ("General wound care",       "30 products"),
        ("Crepe bandage",            "20 products"),
        ("Hydrocolloid",             "8 products"),
        ("Film",                     "8 products"),
        ("Hydrogel",                 "7 products"),
        ("Low adherent",             "7 products"),
        ("Alginate",                 "3 products"),
        ("Iodine",                   "3 products"),
        ("Foam",                     "2 products"),
        ("Hydrofibre / Charcoal",    "1 product each"),
    ], col_widths=(5, 12))

    sub_heading(doc, "Product Gallery Implementation Plan")
    two_col_table(doc, [
        ("__HDR__Component",    "__HDR__Implementation"),
        ("Mapping layer",       "WOUND_TYPE_TO_DRESSINGS dict: WT1-8 → list of dressing types → "
                                "filter wound_products_bigpharmacy.json → return top 3 products"),
        ("API change",          "POST /get_recommendation response adds product_gallery field: "
                                "[{dressing_type, product_name, price_myr, image_url, product_url}]"),
        ("Frontend",            "New 'Available Products' card below recommendation result. "
                                "Shows product image, name, price in MYR, link to purchase."),
        ("FYP evidence panel",  "Evidence Sources panel retained — source numbers cited in text"),
    ])

    doc.add_page_break()

    # ── SECTION 5: WOUND CATEGORY CLASSIFICATION ──────────────────────────────
    section_heading(doc, "5.  FYP2 Secondary Contribution: Wound Category Classification", fill=C_TEAL)

    info_box(doc,
        "Status: CONDITIONAL — depends on labelled wound image dataset secured within first 3 weeks of FYP2.\n"
        "If dataset NOT confirmed by Week 3: drop from FYP2 scope, document as future work.",
        fill=C_ORANGE_BG, border_color=C_ORANGE)

    two_col_table(doc, [
        ("__HDR__Item",          "__HDR__Detail"),
        ("Goal",                 "Classify wound into category: DFU, VLU, Pressure Ulcer, Burn, Skin Tear, "
                                 "Abrasion, Surgical, Laceration, Vascular"),
        ("Why",                  "Ms Saw's request — category-specific metadata filter in ChromaDB retrieval. "
                                 "DFU retrieval should prioritise EWMA/DFU-specific chunks; "
                                 "Burn retrieval should prioritise ANZBA chunks."),
        ("Model approach",       "EfficientNet-B0 or ViT-Small fine-tuned on labelled wound dataset. "
                                 "Classification head → wound_category label → ChromaDB metadata filter axis."),
        ("Dataset needed",       "Minimum: 200–500 labelled images across 5–6 wound categories. "
                                 "Sources: AZH dataset, Medetec, MICCAI wound challenge, or Ms Saw clinical images."),
        ("New ablation (if done)", "R6: Does wound_category metadata filter improve retrieval for category-specific cases?"),
        ("Integration",          "wound_category label feeds into Sub-query A metadata filter: "
                                 "filter by {wound_type: WT3, wound_category: DFU} for diabetic foot infected wounds."),
        ("Risk",                 "If dataset not secured: this item is DROPPED. G4 Multimodal RAG is sufficient FYP2 scope."),
    ])

    doc.add_page_break()

    # ── SECTION 6: DROPPED — CONVERSATIONAL RAG ───────────────────────────────
    section_heading(doc, "6.  Dropped from FYP2: Conversational RAG", fill="7F7F7F")

    info_box(doc,
        "Decision: DROP conversational multi-turn RAG from FYP2 scope.\n"
        "Moved to: Future Work chapter in FYP2 report.",
        fill=C_RED_BG, border_color=C_RED)

    two_col_table(doc, [
        ("__HDR__Reason", "__HDR__Detail"),
        ("Evaluation cost",
         "Requires constructing a new conversational testset: 20–30 sessions x 3–5 turns = 75–150 cases. "
         "Manual construction of ground-truth answers for multi-turn conversations is 3–4 weeks of work alone."),
        ("Evaluation complexity",
         "New metrics needed: Conversational Faithfulness (per-turn FA), Session Coherence (LLM-as-judge), "
         "Turn-level Safety. These require new RAGAS pipeline design and validation."),
        ("Timeline risk",
         "Solo developer, 4–5 months. Building + evaluating conversational RAG would consume 8–10 weeks, "
         "leaving insufficient time for G4 multimodal ablation which is the stronger contribution."),
        ("Better contribution available",
         "G4 Multimodal RAG directly addresses Ms Saw's 'need photo' feedback and extends R5. "
         "It is a more novel and defensible FYP2 contribution than conversational RAG."),
        ("Future work framing",
         "'A conversational multi-turn extension of VerdaSense using LangChain session memory was designed "
         "but not implemented in FYP2 due to scope constraints. The API contract "
         "(POST /chat, session_id, history) is documented for future implementation.'"),
    ])

    sub_heading(doc, "TTS / STT — Also Optional")
    info_box(doc,
        "Voice input (Whisper STT) and output (edge-TTS) are low-priority and will only be implemented "
        "if all other FYP2 components are complete ahead of schedule. "
        "Not a research contribution — a UX feature. Can be demonstrated without formal evaluation.",
        fill=C_GREY_BG, border_color="AAAAAA")

    doc.add_page_break()

    # ── SECTION 7: TIMELINE ───────────────────────────────────────────────────
    section_heading(doc, "7.  FYP2 Timeline (4–5 Months, Solo)")

    two_col_table(doc, [
        ("__HDR__Week",         "__HDR__Deliverable"),
        ("Week 1–2",
         "Bug fixes: RCH metadata filter, debridement prompt injection, testset corrections.\n"
         "UI redesign: patient-friendly language, remove T.I.M.E. rationale section.\n"
         "Decision gate: Wound category dataset — confirm with Ms Saw or drop."),
        ("Week 2–4",
         "Product gallery: dressing type → Big Pharmacy product mapping layer + frontend card.\n"
         "API extension: product_gallery field in POST /get_recommendation response."),
        ("Week 3–5",
         "G4 infrastructure: VLLM captioner integration (Gemini Vision), caption injection "
         "into Sub-query C and assessment_text.\n"
         "Image selection: finalize WT01–WT08 images, validate with Ms Saw review form."),
        ("Week 5–8",
         "G4 ablation: run G4-A through G4-D notebooks (3 runs each).\n"
         "R6 ablation (if wound category dataset secured): wound_category metadata filter.\n"
         "Ms Saw: follow up on G4 review form responses."),
        ("Week 8–10",
         "G4 ablation analysis and write-up.\n"
         "Caption accuracy analysis: T.I.M.E. alignment check (automated) + Ms Saw concordance.\n"
         "[If time]: G4-E VLLM model comparison (Gemini vs GPT-4o as captioner)."),
        ("Week 10–12",
         "Ms Saw final clinician evaluation: 8 cases, Likert 1–5 rating.\n"
         "Category F OOD testset construction (6–8 cases).\n"
         "G5 abstention rate evaluation."),
        ("Week 12–14",
         "FYP2 report writing: literature review, methodology, results, discussion, future work.\n"
         "[If time]: TTS/STT integration demo."),
        ("Week 15 (buffer)",
         "Final demo preparation, report submission, viva preparation."),
    ])

    doc.add_page_break()

    # ── SECTION 8: RISKS ──────────────────────────────────────────────────────
    section_heading(doc, "8.  Risks & Mitigations", fill=C_ORANGE)

    two_col_table(doc, [
        ("__HDR__Risk",                         "__HDR__Mitigation"),
        ("Wound category dataset not secured",
         "Hard deadline Week 3. If not confirmed → drop entirely. G4 is sufficient FYP2 scope."),
        ("VLLM infection detection failure\n(WT03/04 issue from R5)",
         "T.I.M.E. infection label retains priority in conflict. "
         "VLLM enriches tissue/moisture/edge, not infection axis. "
         "Document as finding: 'Early local infection not visually detectable by VLLM.'"),
        ("Ms Saw unavailable for review",
         "G4 review form is designed for async completion. "
         "WhatsApp-based — can be completed at her convenience. "
         "Even partial responses (image suitability only) are usable."),
        ("API cost (G4 ablation runs)",
         "G4 has 5 configs x 3 runs = 15 RAGAS evaluation runs. "
         "Estimated cost: RM 30–50 total. Acceptable."),
        ("G4 results show no improvement\nover G2-D baseline",
         "Null result is still publishable for FYP. "
         "Finding: 'VLLM caption injection does not significantly improve FA/Safety for "
         "structured T.I.M.E. inputs.' "
         "Supports future work on better fusion strategies."),
    ])

    doc.add_page_break()

    # ── SECTION 9: QUESTIONS ──────────────────────────────────────────────────
    section_heading(doc, "9.  Questions for Supervisor")

    questions = [
        ("Q1 — G4 Scope",
         "Is G4 multimodal ablation (G4-A through G4-D/E) an academically sufficient primary FYP2 "
         "contribution, given R5 already covers the retrieval half?"),
        ("Q2 — Wound Category Classification",
         "If Ms Saw cannot provide a labelled wound image dataset within 3 weeks, is it acceptable "
         "to fully drop Proposal A (wound category classification) from FYP2 and document it as future work?"),
        ("Q3 — Caption-T.I.M.E. Fusion",
         "The caption-T.I.M.E. conflict handling (T.I.M.E. priority for infection axis) is currently "
         "rule-based. Is a rule-based fusion sufficient for FYP2, or should a learned fusion layer "
         "be included?"),
        ("Q4 — Human Evaluation",
         "Ms Saw's 8-image review + final Likert evaluation: is this clinically sufficient for "
         "the 'human evaluation' chapter in FYP2, or is a larger panel of reviewers needed?"),
        ("Q5 — Report Structure",
         "Should FYP2 report Chapter 4 (Results) present FYP1 ablation results (R1-R5, G1-G3) "
         "briefly as 'baseline' and focus Chapter 4 on G4 multimodal results, "
         "or should all ablation results have equal weight?"),
    ]

    for qnum, (title, body) in enumerate(questions, 1):
        qt = doc.add_table(rows=1, cols=1)
        qt.style = "Table Grid"
        c = qt.rows[0].cells[0]
        shade_cell(c, C_LIGHT_BLUE)
        cell_border(c, C_MED_BLUE, "8")
        c.paragraphs[0].clear()
        tp = c.add_paragraph(title)
        tp.runs[0].bold = True; tp.runs[0].font.size = Pt(10)
        tp.runs[0].font.color.rgb = RGBColor.from_string(C_DARK_BLUE)
        bp = c.add_paragraph(body)
        bp.runs[0].font.size = Pt(9)
        bp.paragraph_format.space_before = Pt(3)
        bp.paragraph_format.space_after  = Pt(3)
        resp = c.add_paragraph("Supervisor response: ___________________________________________"
                               "_____________________________________")
        resp.runs[0].font.size = Pt(9)
        resp.runs[0].italic = True
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph("VerdaSense FYP2 Supervisor Meeting Proposal · "
                           "Tee Qi Jing (23004894) · Universiti Malaya · June 2026")
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUT)
    print(f"[OK] Saved: {OUT}")

if __name__ == "__main__":
    build()
