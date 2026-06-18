"""
VerdaSense G4 Multimodal — Ms Saw Clinical Review Form Generator
Generates: MDs/FYP2 Migration/VerdaSense_G4_Clinical_Review_Form.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE = r"C:\Users\GIGA\OneDrive - Universiti Malaya\Documents\rag-for-beginners"
IMG_DIR = os.path.join(BASE, "ragas_testset", "wound_images")
OUT_PATH = os.path.join(BASE, "MDs", "FYP2 Migration",
                        "VerdaSense_G4_Clinical_Review_Form.docx")

# ── Wound type data ────────────────────────────────────────────────────────────
WOUND_TYPES = [
    {
        "wt": "01", "label": "Type 1 — Clean Granulating, Dry, Not Infected",
        "image": "WT01_wound.jpeg",
        "time": {"T": "100% Granulation", "I": "Not infected",
                 "M": "Low to moderate exudate", "E": "Advancing"},
        "abx": False, "referral": False,
        "linked_cases": "cat_a_type1_dry · cat_b_silver_clean_granulating · cat_b_skin_tear_fragile "
                        "· cat_b_postop_clean · cat_b_burns_hand · cat_b_skin_tear_type2_flap "
                        "· cat_b_burns_minor_epidermal · cat_c_film_vs_hydrocolloid",
        "caption": (
            "TISSUE: Approximately 90% healthy red granulation tissue fills the wound bed, "
            "with 10% pale pink epithelial tissue advancing from the edges.\n"
            "INFECTION: No visible infection signs.\n"
            "MOISTURE: The wound bed appears moist without excessive exudate. "
            "Periwound skin is dry, showing no maceration.\n"
            "EDGE: Wound edges exhibit slight epibole, with some epithelialization noted. "
            "Edges are generally well-defined but show minor undermining.\n"
            "ADDITIONAL: The wound measures approximately 1 cm × 1 cm, shallow depth. "
            "Periwound skin displays mild erythema, likely due to inflammation rather than infection."
        ),
        "verdaSense_reco": "Primary: Film dressing (Alt: Thin hydrocolloid) — change every 2–5 days",
        "contraindicated": "Silver dressings · Charcoal dressings",
        "panel": "Panel 1: Agreed | Panel 2: Agreed | Panel 3: Agreed",
    },
    {
        "wt": "02", "label": "Type 2 — Granulating, High Exudate, Not Infected",
        "image": "WT02_medetec_0116.png",
        "time": {"T": "90–95% Granulation, 5–10% NV (slough)",
                 "I": "Not infected", "M": "High exudate", "E": "Advancing"},
        "abx": False, "referral": False,
        "linked_cases": "cat_a_type2_wet · cat_c_dressing_saturation · cat_c_heavy_exudate_maceration "
                        "· cat_d_notes_malodour_clean",
        "caption": (
            "TISSUE: Approximately 60% healthy red granulation tissue, 30% yellow fibrinous slough, "
            "10% pale avascular tissue.\n"
            "INFECTION: No visible infection signs.\n"
            "MOISTURE: Moderate exudate; wound bed appears moist with serous fluid. "
            "Periwound skin is slightly macerated.\n"
            "EDGE: Wound edges show epibole with some undermining noted. "
            "Epithelialization is minimal, edges appear non-advancing.\n"
            "ADDITIONAL: Estimated wound size approximately 2 cm × 2 cm. "
            "Depth appears superficial. Periwound skin exhibits mild erythema and dryness."
        ),
        "verdaSense_reco": "Primary: Alginate (sheet/rope) (Alt: Hydrofibre) — "
                           "change every 2–5 days; secondary dressing required",
        "contraindicated": "None specified by algorithm",
        "panel": "Panel 1: Agreed | Panel 2: Agreed | Panel 3: Agreed",
    },
    {
        "wt": "03", "label": "Type 3 — Locally Infected, Low Exudate",
        "image": "WT03_wound.jpeg",
        "time": {"T": "~75% Granulation, ~25% NV (slough+necrosis)",
                 "I": "Locally infected", "M": "Low to moderate exudate", "E": "Non-advancing"},
        "abx": True, "referral": False,
        "linked_cases": "cat_a_type3_dry_infected · cat_b_iodine_thyroid · cat_b_diabetic_foot "
                        "· cat_c_dry_infected_combo · cat_d_notes_diabetic_nonhealing · cat_e_dfu_infected_ewma",
        "caption": (
            "TISSUE: Approximately 80% of the wound bed covered with healthy pink granulation tissue; "
            "20% shows pale, edematous tissue at the periphery.\n"
            "INFECTION: No visible infection signs. [NOTE: T.I.M.E. profile indicates locally infected — "
            "image may not visually show all infection signs]\n"
            "MOISTURE: The wound bed appears moist with a small amount of serous exudate. "
            "Periwound skin is intact without maceration.\n"
            "EDGE: The wound edges exhibit epibole, with rolled tissue overhanging the wound bed. "
            "No evidence of epithelialization.\n"
            "ADDITIONAL: The wound measures approximately 2 cm × 1.5 cm. "
            "Depth is shallow. Periwound skin is slightly erythematous."
        ),
        "verdaSense_reco": "Primary: Silver dressing (Alt: Hydrogel with secondary dressing) "
                           "— change every 2–3 days",
        "contraindicated": "None specified by algorithm",
        "panel": (
            "Panel 1: Debridement of edges/slough with daily hydrogel; if 2–3 days, preferably "
            "hydrofibre + ABx\n"
            "Panel 2: Modern dressing and assess wound progression\n"
            "Panel 3: Agreed + wound debridement of the slough"
        ),
    },
    {
        "wt": "04", "label": "Type 4 — Locally Infected, High Exudate",
        "image": "WT04_wsnet_0816.png",
        "time": {"T": "~80% Granulation, 12–20% NV (slough)", "I": "Locally infected",
                 "M": "High exudate", "E": "Non-advancing"},
        "abx": True, "referral": False,
        "linked_cases": "cat_a_type4_wet_infected · cat_d_notes_infection_override",
        "caption": (
            "TISSUE: Approximately 85% healthy red granulation tissue, 15% pale yellow fibrinous slough "
            "along the edges. No necrotic tissue observed.\n"
            "INFECTION: No visible infection signs. [NOTE: T.I.M.E. profile indicates locally infected]\n"
            "MOISTURE: The wound bed appears moist with a thin layer of serous exudate. "
            "Periwound skin shows no maceration.\n"
            "EDGE: Wound edges exhibit slight epibole, appearing rolled. Edges are undermined slightly.\n"
            "ADDITIONAL: The wound measures approximately 4 cm × 3 cm. "
            "Depth appears shallow, less than 0.5 cm. Periwound skin intact but mildly erythematous."
        ),
        "verdaSense_reco": "Primary: Silver dressing + Alginate (Alt: Silver + Hydrofibre) "
                           "— change every 2–3 days; secondary dressing required",
        "contraindicated": "None specified by algorithm",
        "panel": (
            "Panel 1: Debridement with the dressing + ABx\n"
            "Panel 2: Suggest daily dressing instead of modern dressing as high exudate "
            "+/- bedside deslough first\n"
            "Panel 3: Agreed + wound debridement of the slough"
        ),
    },
    {
        "wt": "05", "label": "Type 5 — Necrotic/Sloughy, Dry, Not Infected",
        "image": "WT05_wsnet_0384.png",
        "time": {"T": "40–50% Necrotic/Slough, ~30% Granulation",
                 "I": "Not infected", "M": "Low/dry exudate", "E": "Non-advancing"},
        "abx": False, "referral": False,
        "linked_cases": "cat_a_type5_dry_necrotic · cat_b_alginate_dry_wound · cat_b_honey_dry_necrotic",
        "caption": (
            "TISSUE: Approximately 50% healthy red granulation tissue, 30% yellow fibrinous slough, "
            "20% dark necrotic eschar at the base.\n"
            "INFECTION: No visible infection signs.\n"
            "MOISTURE: The wound bed appears moist without excessive exudate. "
            "Periwound skin is intact but slightly erythematous.\n"
            "EDGE: Wound edges show rolled epibole with some undermining noted. "
            "Epithelialization is minimal, edges appear non-advancing.\n"
            "ADDITIONAL: Estimated wound size approximately 2 cm × 2 cm. "
            "Depth varies with some areas of deeper tissue loss. "
            "Periwound skin shows mild erythema."
        ),
        "verdaSense_reco": "Primary: Hydrogel (no direct alternative) — change every 2–3 days; "
                           "secondary dressing required",
        "contraindicated": "Alginate · Hydrofibre (not suitable for dry wounds)",
        "panel": (
            "Panel 1: Debridement and daily hydrogel; need expert opinion and assess vascular supply "
            "(sometimes associated arterial compromise)\n"
            "Panel 2: Chemical debridement then bedside deslough\n"
            "Panel 3: Agreed"
        ),
    },
    {
        "wt": "06", "label": "Type 6 — Necrotic/Sloughy, High Exudate, Not Infected",
        "image": "WT06_medetec_0298.png",
        "time": {"T": "30–35% Necrotic, 30–35% Slough, ~35% Granulation",
                 "I": "Not infected", "M": "High exudate", "E": "Non-advancing"},
        "abx": False, "referral": True,
        "linked_cases": "cat_a_type6_wet_necrotic · cat_b_referral_type6 · cat_d_notes_npwt_adjunct "
                        "· cat_e_vlu_chronic_ewma",
        "caption": (
            "TISSUE: Approximately 50% healthy red granulation tissue, 40% pale yellow fibrinous slough, "
            "10% dark necrotic eschar at the base.\n"
            "INFECTION: No visible infection signs.\n"
            "MOISTURE: Moderate exudate; wound bed appears moist. "
            "Periwound skin is slightly macerated due to moisture exposure.\n"
            "EDGE: Wound edges show epibole with rolled appearance, undermining noted along lateral aspect. "
            "Epithelialization is minimal.\n"
            "ADDITIONAL: Estimated wound size approximately 2 cm × 3 cm. "
            "Periwound skin exhibits mild erythema, possibly due to moisture."
        ),
        "verdaSense_reco": "Primary: Alginate (sheet/rope) (Alt: Hydrofibre) — "
                           "change every 2–5 days; referral to specialist indicated",
        "contraindicated": "None specified by algorithm",
        "panel": (
            "Panel 1: Debridement with daily or bd Dermasyn dressing for 1–2 days; "
            "when debridement assessment no longer required, change to Ag Alginate\n"
            "Panel 2: Daily dressing + deslough\n"
            "Panel 3: Agreed + wound debridement of the slough"
        ),
    },
    {
        "wt": "07", "label": "Type 7 — Infected + Necrotic, Low Exudate (Complex)",
        "image": "WT07_wsnet_0539.png",
        "time": {"T": "25–40% Necrotic, 25–30% Slough, ~20% Granulation",
                 "I": "Locally infected", "M": "Low to moderate exudate", "E": "Non-advancing"},
        "abx": True, "referral": True,
        "linked_cases": "cat_a_type7_dry_infected_necrotic · cat_c_time_assessment_mixed",
        "caption": (
            "TISSUE: Approximately 50% dark necrotic eschar centrally, 30% yellow fibrinous slough "
            "along edges, 20% granulation tissue visible beneath slough.\n"
            "INFECTION: Visible infection indicators — perilesional erythema, localised warmth, "
            "and purulent exudate suggest active infection.\n"
            "MOISTURE: Moderate exudate level; wound bed appears moist with visible cloudy discharge. "
            "Periwound skin shows early signs of maceration.\n"
            "EDGE: Wound edges exhibit epibole with rolled appearance, undermining noted at lower margin. "
            "Epithelialization is absent; edges appear non-advancing.\n"
            "ADDITIONAL: Estimated wound size approximately 5 cm × 4 cm. "
            "Periwound skin displays mild erythema and oedema."
        ),
        "verdaSense_reco": "Primary: Silver dressing + Hydrogel (hydrogel for autolytic debridement, "
                           "silver for infection) — change every 2–3 days; specialist referral required",
        "contraindicated": "Alginate · Hydrofibre (not suitable for dry/infected necrotic wounds)",
        "panel": (
            "Panel 1: Debridement with daily or bd Dermasyn dressing until debridement no longer required, "
            "then change to Ag Alginate + ABx\n"
            "Panel 2: Daily dressing + deslough + antibiotic; might need OT debridement\n"
            "Panel 3: Wound debridement then dressing + antibiotic"
        ),
    },
    {
        "wt": "08", "label": "Type 8 — Infected + Necrotic, High Exudate (Most Complex)",
        "image": "WT08_medetec_0175.png",
        "time": {"T": "30–50% Necrotic, 20–35% Slough, ~35% Granulation",
                 "I": "Locally infected", "M": "High exudate", "E": "Non-advancing"},
        "abx": True, "referral": True,
        "linked_cases": "cat_a_type8_wet_infected_necrotic · cat_b_npwt_necrotic_eschar "
                        "· cat_c_malodour_type8",
        "caption": (
            "TISSUE: Approximately 40% healthy red granulation tissue, 50% greenish-yellow slough "
            "covering the wound bed, and 10% dark necrotic tissue at the base.\n"
            "INFECTION: Visible infection indicators — green/cloudy discharge, perilesional erythema, "
            "and possible purulent exudate.\n"
            "MOISTURE: The wound bed is moist with visible exudate that appears cloudy and purulent. "
            "Periwound skin shows signs of maceration.\n"
            "EDGE: The wound edges exhibit epibole with some undermining noted. "
            "Edges appear rolled and non-advancing, lacking epithelialization.\n"
            "ADDITIONAL: The wound measures approximately 5 cm × 3 cm with depth suggesting involvement "
            "beyond superficial layers. Periwound skin is erythematous and macerated."
        ),
        "verdaSense_reco": "Primary: Silver dressing + Alginate (Alt: Silver + Hydrofibre; "
                           "charcoal if malodour) — change every 2–3 days; urgent specialist referral",
        "contraindicated": "None specified by algorithm",
        "panel": (
            "Panel 1: Debridement with daily or bd Dermasyn dressing until debridement no longer "
            "required, then change to Ag Alginate + ABx\n"
            "Panel 2: Daily dressing + deslough + antibiotic; might need OT debridement\n"
            "Panel 3: Wound debridement then dressing + antibiotic"
        ),
    },
]

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top="single", bottom="single", left="single", right="single", sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "444444")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_colored_para(cell, text, bold=False, italic=False,
                     font_size=10, color="000000", space_before=0, space_after=0):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_label_value_row(table, label, value, label_bg="EAF0FB", value_bg="FFFFFF"):
    row = table.add_row()
    lc = row.cells[0]
    vc = row.cells[1]
    set_cell_bg(lc, label_bg)
    set_cell_bg(vc, value_bg)
    set_cell_borders(lc)
    set_cell_borders(vc)
    lc.paragraphs[0].clear()
    vc.paragraphs[0].clear()
    add_colored_para(lc, label, bold=True, font_size=9, color="1F3864")
    add_colored_para(vc, value, font_size=9)


def add_fillable_row(table, label, hint="", label_bg="EAF0FB", rows=2):
    row = table.add_row()
    lc = row.cells[0]
    vc = row.cells[1]
    set_cell_bg(lc, label_bg)
    set_cell_bg(vc, "FFFEF0")   # light yellow = fillable
    set_cell_borders(lc)
    set_cell_borders(vc)
    lc.paragraphs[0].clear()
    vc.paragraphs[0].clear()
    add_colored_para(lc, label, bold=True, font_size=9, color="1F3864")
    if hint:
        add_colored_para(vc, hint, italic=True, font_size=8, color="888888")
    # Add empty lines for writing space
    for _ in range(rows):
        add_colored_para(vc, "", font_size=9)


def add_checkbox_row(table, label, options, label_bg="EAF0FB"):
    row = table.add_row()
    lc = row.cells[0]
    vc = row.cells[1]
    set_cell_bg(lc, label_bg)
    set_cell_bg(vc, "FFFEF0")
    set_cell_borders(lc)
    set_cell_borders(vc)
    lc.paragraphs[0].clear()
    vc.paragraphs[0].clear()
    add_colored_para(lc, label, bold=True, font_size=9, color="1F3864")
    checkbox_text = "    ".join(f"☐  {opt}" for opt in options)
    add_colored_para(vc, checkbox_text, font_size=9)
    add_colored_para(vc, "Comments / Corrections:", italic=True, font_size=8, color="888888")
    add_colored_para(vc, "", font_size=9)
    add_colored_para(vc, "", font_size=9)


# ── Main document builder ──────────────────────────────────────────────────────

def build_form():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Cover page ─────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("VerdaSense — G4 Multimodal Experiment")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Clinical Image Review Form")
    r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    meta_lines = [
        ("Reviewer", "Ms Saw Shier Khee, Surgical Collaborator"),
        ("Student",  "Tee Qi Jing (23004894)"),
        ("Project",  "VerdaSense RAG — FYP2, Universiti Malaya"),
        ("Date",     "_____________________________  2026"),
    ]
    for label, val in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  ")
        r.bold = True; r.font.size = Pt(11)
        r2 = p.add_run(val)
        r2.font.size = Pt(11)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    r = intro.add_run("Purpose of This Form")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    doc.add_paragraph(
        "VerdaSense FYP2 is extending the wound dressing RAG system to accept actual wound photographs "
        "as additional input. A Vision AI model (VLLM) generates a clinical description from each image, "
        "which is then used to enrich the dressing recommendation.\n\n"
        "This form asks you to review 8 wound images (one per wound type, WT01–WT08) and provide:\n"
        "  1. Image suitability — is this a clinically reasonable example of the wound type?\n"
        "  2. AI caption accuracy — is the VLLM's clinical description accurate?\n"
        "  3. Your clinical recommendation — what you would recommend for this wound.\n\n"
        "Your responses will become the GOLD STANDARD for the G4 multimodal evaluation. "
        "Yellow-shaded cells are for you to fill in. No cell is mandatory — any input is valuable.",
        style="Normal"
    ).runs[0].font.size = Pt(10)

    # Colour legend
    doc.add_paragraph()
    leg = doc.add_paragraph()
    r = leg.add_run("Colour Legend:  ")
    r.bold = True; r.font.size = Pt(9)
    leg.add_run("Blue header = wound info (read only)    ").font.size = Pt(9)
    leg.add_run("Yellow cell = please fill in    ").font.size = Pt(9)
    leg.add_run("Grey cell = reference / system output (read only)").font.size = Pt(9)

    doc.add_page_break()

    # ── One section per wound type ─────────────────────────────────────────────
    for wt_data in WOUND_TYPES:
        wt     = wt_data["wt"]
        label  = wt_data["label"]
        abx    = wt_data["abx"]
        ref    = wt_data["referral"]

        # Section heading
        h = doc.add_paragraph()
        r = h.add_run(f"  WT{wt} — {label}  ")
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after  = Pt(4)
        shd = OxmlElement("w:pPr")
        h._p.insert(0, shd)
        # shade the heading paragraph background via direct XML
        pPr = h._p.get_or_add_pPr()
        pShd = OxmlElement("w:shd")
        pShd.set(qn("w:val"), "clear")
        pShd.set(qn("w:color"), "auto")
        pShd.set(qn("w:fill"), "1F3864")
        pPr.append(pShd)

        # Flags row
        flags = []
        if abx: flags.append("⚠  Antibiotic consideration required")
        if ref: flags.append("⬆  Specialist referral indicated")
        if not abx and not ref: flags.append("✓  No antibiotic · No referral required")
        fp = doc.add_paragraph("    ".join(flags))
        fp.runs[0].font.size = Pt(9)
        fp.runs[0].bold = True
        fp.runs[0].font.color.rgb = RGBColor(0xC5, 0x50, 0x00) if (abx or ref) else RGBColor(0x37, 0x63, 0x31)
        fp.paragraph_format.space_before = Pt(2)
        fp.paragraph_format.space_after  = Pt(4)

        # ── Top table: image + T.I.M.E. profile ───────────────────────────────
        top_tbl = doc.add_table(rows=1, cols=2)
        top_tbl.style = "Table Grid"
        top_tbl.columns[0].width = Cm(8)
        top_tbl.columns[1].width = Cm(9)
        img_cell  = top_tbl.rows[0].cells[0]
        time_cell = top_tbl.rows[0].cells[1]
        set_cell_bg(img_cell,  "F0F4F8")
        set_cell_bg(time_cell, "EAF0FB")
        set_cell_borders(img_cell)
        set_cell_borders(time_cell)

        # Insert image
        img_path = os.path.join(IMG_DIR, wt_data["image"])
        img_cell.paragraphs[0].clear()
        p_img = img_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(img_path):
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(2.8))
        else:
            p_img.add_run(f"[Image not found: {wt_data['image']}]").font.size = Pt(9)
        # Image caption below
        cap_p = img_cell.add_paragraph(wt_data["image"])
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.runs[0].font.size = Pt(7)
        cap_p.runs[0].italic = True
        cap_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # T.I.M.E. profile
        time_cell.paragraphs[0].clear()
        th = time_cell.add_paragraph("T.I.M.E. Wound Profile")
        th.runs[0].bold = True
        th.runs[0].font.size = Pt(10)
        th.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        th.paragraph_format.space_after = Pt(4)

        time_labels = {"T": "Tissue", "I": "Infection", "M": "Moisture", "E": "Edge"}
        for key, desc in wt_data["time"].items():
            tp = time_cell.add_paragraph()
            r1 = tp.add_run(f"{key} ({time_labels[key]}):  ")
            r1.bold = True; r1.font.size = Pt(9)
            r2 = tp.add_run(desc)
            r2.font.size = Pt(9)
            tp.paragraph_format.space_before = Pt(1)
            tp.paragraph_format.space_after  = Pt(1)

        # Linked cases
        time_cell.add_paragraph()
        lc_p = time_cell.add_paragraph()
        r1 = lc_p.add_run("Linked test cases:  ")
        r1.bold = True; r1.font.size = Pt(8)
        r2 = lc_p.add_run(wt_data["linked_cases"])
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ── Main review table ─────────────────────────────────────────────────
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Cm(5)
        tbl.columns[1].width = Cm(12)

        # Column headers
        hrow = tbl.rows[0]
        set_cell_bg(hrow.cells[0], "2E74B5")
        set_cell_bg(hrow.cells[1], "2E74B5")
        set_cell_borders(hrow.cells[0])
        set_cell_borders(hrow.cells[1])
        hrow.cells[0].paragraphs[0].clear()
        hrow.cells[1].paragraphs[0].clear()
        r = hrow.cells[0].add_paragraph("Field").runs[0]
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r = hrow.cells[1].add_paragraph("Content / Response").runs[0]
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Row 1: AI-generated caption (read-only, grey)
        row = tbl.add_row()
        set_cell_bg(row.cells[0], "EAF0FB")
        set_cell_bg(row.cells[1], "F5F5F5")
        set_cell_borders(row.cells[0])
        set_cell_borders(row.cells[1])
        row.cells[0].paragraphs[0].clear()
        row.cells[1].paragraphs[0].clear()
        r = row.cells[0].add_paragraph("AI Vision Caption\n(VLLM output\n— read only)")
        r.runs[0].bold = True; r.runs[0].font.size = Pt(9); r.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        cap_para = row.cells[1].add_paragraph(wt_data["caption"])
        cap_para.runs[0].font.size = Pt(8)
        cap_para.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        cap_para.runs[0].italic = True

        # Row 2: Image suitability checkbox
        add_checkbox_row(tbl, "1. Image Suitability\n(Does this image\nrepresent this\nwound type?)",
                         ["Suitable — image is a reasonable example",
                          "Not suitable — image is misleading / incorrect"])

        # Row 3: Caption accuracy checkbox
        add_checkbox_row(tbl, "2. AI Caption\nAccuracy\n(Is the VLLM\ndescription\nclinically correct?)",
                         ["Accurate — VLLM description is clinically sound",
                          "Partially accurate — some errors (note below)",
                          "Inaccurate — description is misleading"])

        # Row 4: Caption corrections
        add_fillable_row(tbl, "3. Caption\nCorrections\n(What did the AI\nmiss or get wrong?)",
                         hint="e.g. 'Infection signs visible but VLLM missed them' or 'Tissue % estimate is off'",
                         rows=2)

        # Row 5: Clinical description from image
        add_fillable_row(tbl, "4. Your Clinical\nDescription\n(What do you see\nin this image?)",
                         hint="Describe tissue, infection signs, exudate, wound edges — as you would document it clinically",
                         rows=3)

        # Row 6: Primary dressing
        add_fillable_row(tbl, "5. Primary Dressing\nRecommendation\n(Based on image +\nwound type)",
                         hint="e.g. 'Silver dressing, change every 2–3 days' or 'Hydrogel + secondary foam'",
                         rows=2)

        # Row 7: Secondary dressing
        add_fillable_row(tbl, "6. Secondary\nDressing\n(If applicable)",
                         hint="e.g. 'Foam as secondary' or 'Not required'",
                         rows=1)

        # Row 8: Debridement
        add_checkbox_row(tbl, "7. Debridement\nRequired?",
                         ["Not required", "Autolytic (dressing-based)",
                          "Bedside deslough", "Surgical / OT debridement"])

        # Row 9: Antibiotic
        add_checkbox_row(tbl, "8. Antibiotic\nRecommendation",
                         ["Antibiotic NOT indicated",
                          "Topical antimicrobial dressing only (no systemic ABx)",
                          "Systemic antibiotic recommended"])

        # Row 10: Referral
        add_checkbox_row(tbl, "9. Referral /\nEscalation",
                         ["No referral — manageable in community",
                          "Referral recommended to GP / clinic",
                          "Urgent specialist referral required"])

        # Row 11: Additional notes
        add_fillable_row(tbl, "10. Additional\nClinical Notes\n(Anything else?)",
                         hint="Any safety concerns, contraindications, Malaysian-specific considerations, etc.",
                         rows=3)

        # ── Reference row: VerdaSense system output ────────────────────────────
        ref_row = tbl.add_row()
        set_cell_bg(ref_row.cells[0], "D6E4F0")
        set_cell_bg(ref_row.cells[1], "EBF5FB")
        set_cell_borders(ref_row.cells[0])
        set_cell_borders(ref_row.cells[1])
        ref_row.cells[0].paragraphs[0].clear()
        ref_row.cells[1].paragraphs[0].clear()
        r = ref_row.cells[0].add_paragraph("VerdaSense\nSystem Output\n(for comparison\n— read only)")
        r.runs[0].bold = True; r.runs[0].font.size = Pt(9)
        r.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        vs_p = ref_row.cells[1].add_paragraph(f"Recommendation: {wt_data['verdaSense_reco']}\n"
                                               f"Contraindicated: {wt_data['contraindicated']}")
        vs_p.runs[0].font.size = Pt(8)

        # ── Prior panel comments ───────────────────────────────────────────────
        panel_row = tbl.add_row()
        set_cell_bg(panel_row.cells[0], "D6E4F0")
        set_cell_bg(panel_row.cells[1], "EBF5FB")
        set_cell_borders(panel_row.cells[0])
        set_cell_borders(panel_row.cells[1])
        panel_row.cells[0].paragraphs[0].clear()
        panel_row.cells[1].paragraphs[0].clear()
        r = panel_row.cells[0].add_paragraph("Prior Panel\nComments\n(from Cat A\nreview — read only)")
        r.runs[0].bold = True; r.runs[0].font.size = Pt(9)
        r.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        pp = panel_row.cells[1].add_paragraph(wt_data["panel"])
        pp.runs[0].font.size = Pt(8)
        pp.runs[0].italic = True

        doc.add_page_break()

    # ── Overall comments page ──────────────────────────────────────────────────
    h = doc.add_paragraph("Overall Comments & Summary")
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    h.paragraph_format.space_after = Pt(8)

    oc_tbl = doc.add_table(rows=1, cols=2)
    oc_tbl.style = "Table Grid"
    oc_tbl.columns[0].width = Cm(5)
    oc_tbl.columns[1].width = Cm(12)

    oc_items = [
        ("Images accepted\n(WT numbers)",
         "Which of the 8 images are suitable? List WT numbers:  _______________________"),
        ("Images rejected\n(WT numbers + reason)",
         "Which images need replacing and why?"),
        ("Missing wound\ntypes / scenarios",
         "Are there wound types important for Malaysian patients NOT covered by WT01–WT08?"),
        ("VLLM caption\noverall quality",
         "Overall, how well does the AI vision model describe wound characteristics?\n"
         "☐  Good — clinically useful descriptions\n"
         "☐  Partial — some useful, some inaccurate\n"
         "☐  Poor — descriptions cannot be trusted clinically"),
        ("Recommended\nchanges to\nVerdaSense system",
         "What is the most important clinical improvement you would suggest for FYP2?"),
        ("Availability of\nclinical images\n(optional)",
         "Do you have access to de-identified wound photographs from clinical practice?\n"
         "☐  Yes — I can share (please advise on ethics/consent process)\n"
         "☐  No — public datasets are sufficient for this research purpose"),
        ("Willingness to\nconduct final\nhuman evaluation\n(15–20 min, online)",
         "At project completion, would you be willing to rate 8–10 AI recommendations\n"
         "on clinical accuracy (Likert 1–5 scale)?\n"
         "☐  Yes  ☐  No  ☐  Maybe — contact me closer to date"),
    ]

    hrow = oc_tbl.rows[0]
    set_cell_bg(hrow.cells[0], "1F3864")
    set_cell_bg(hrow.cells[1], "1F3864")
    hrow.cells[0].paragraphs[0].clear()
    hrow.cells[1].paragraphs[0].clear()
    r = hrow.cells[0].add_paragraph("Question").runs[0]
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r = hrow.cells[1].add_paragraph("Your Response").runs[0]
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for label, hint in oc_items:
        row = oc_tbl.add_row()
        set_cell_bg(row.cells[0], "EAF0FB")
        set_cell_bg(row.cells[1], "FFFEF0")
        set_cell_borders(row.cells[0])
        set_cell_borders(row.cells[1])
        row.cells[0].paragraphs[0].clear()
        row.cells[1].paragraphs[0].clear()
        rp = row.cells[0].add_paragraph(label)
        rp.runs[0].bold = True; rp.runs[0].font.size = Pt(9)
        rp.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        vp = row.cells[1].add_paragraph(hint)
        vp.runs[0].font.size = Pt(9)
        row.cells[1].add_paragraph("")

    # Footer note
    doc.add_paragraph()
    fn = doc.add_paragraph(
        "Thank you for your time, Ms Saw. Your clinical expertise directly improves the accuracy and "
        "safety of VerdaSense. Please return this form via WhatsApp or email to Tee Qi Jing (23004894)."
    )
    fn.runs[0].font.size = Pt(9)
    fn.runs[0].italic = True

    doc.add_paragraph()
    sig = doc.add_paragraph("Prepared by Tee Qi Jing (23004894) · VerdaSense FYP2 · "
                            "Universiti Malaya · June 2026")
    sig.runs[0].font.size = Pt(8)
    sig.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUT_PATH)
    print(f"[OK] Saved: {OUT_PATH}")


if __name__ == "__main__":
    build_form()
